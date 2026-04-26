from docx import Document
from pathlib import Path
import os
import re
from eva.utils.logger import get_logger

logger = get_logger("EVA.actions.documents")

def add_heading(doc: Document, text: str, level: int):
    """Helper to add heading."""
    doc.add_heading(text, level=level)

def add_paragraph(doc: Document, text: str):
    """Helper to add paragraph."""
    doc.add_paragraph(text)

def save_document(doc: Document, filename: str) -> str:
    """Saves to desktop, returns file path."""
    profile = Path(os.environ['USERPROFILE'])
    desktop = profile / 'OneDrive' / 'Desktop'
    if not desktop.exists():
        desktop = profile / 'Desktop'
    
    filepath = desktop / filename
    doc.save(filepath)
    return str(filepath)

def convert_to_pdf(docx_path: str):
    """Safely converts docx to pdf."""
    try:
        from docx2pdf import convert
        pdf_path = docx_path.replace(".docx", ".pdf")
        convert(docx_path, pdf_path)
        logger.info(f"Converted to PDF: {pdf_path}")
    except ImportError:
        logger.warning("docx2pdf not installed. Skipping PDF conversion.")
    except Exception as e:
        logger.warning(f"PDF conversion failed: {e}")

def create_srs_document(project_name: str, content: dict, preview: bool = False) -> str:
    """Creates full SRS document."""
    try:
        doc = Document()
        doc.add_heading(f'Software Requirements Specification: {project_name}', 0)
        
        # Introduction
        add_heading(doc, '1. Introduction', level=1)
        add_paragraph(doc, content.get("introduction", "This document outlines the SRS for the system."))
        
        # Scope
        add_heading(doc, '2. Scope', level=1)
        add_paragraph(doc, content.get("scope", "The project scope encompasses all features listed below."))
        
        # Functional Requirements
        add_heading(doc, '3. Functional Requirements', level=1)
        for req in content.get("functional", ["Requirement 1", "Requirement 2"]):
            doc.add_paragraph(req, style='List Bullet')
            
        # Non-Functional Requirements
        add_heading(doc, '4. Non-Functional Requirements', level=1)
        for req in content.get("non_functional", ["Performance constraints", "Security policy"]):
            doc.add_paragraph(req, style='List Bullet')
            
        # System Architecture
        add_heading(doc, '5. System Architecture', level=1)
        add_paragraph(doc, content.get("architecture", "A high-level client-server architecture description."))
        
        safe_name = re.sub(r'[^a-zA-Z0-9_\-]', '_', project_name)
        filename = f"{safe_name}_SRS.docx"
        
        filepath = save_document(doc, filename)
        logger.info(f"Created SRS at {filepath}")
        
        convert_to_pdf(filepath)
        if preview:
            os.startfile(filepath)
            
        return filepath
    except Exception as e:
        logger.error(f"SRS generation error: {e}")
        return ""

def create_report(title: str, sections: list[dict], preview: bool = False) -> str:
    """Generic report generator. sections = [{'heading': '', 'content': ''}]"""
    try:
        doc = Document()
        doc.add_heading(title, 0)
        
        for idx, sec in enumerate(sections):
            add_heading(doc, f"{idx+1}. {sec.get('heading', '')}", level=1)
            add_paragraph(doc, sec.get('content', ''))
            
        safe_name = re.sub(r'[^a-zA-Z0-9_\-]', '_', title)
        filename = f"{safe_name}_Report.docx"
        filepath = save_document(doc, filename)
        
        convert_to_pdf(filepath)
        if preview:
            os.startfile(filepath)
            
        return filepath
    except Exception as e:
        logger.error(f"Report generation error: {e}")
        return ""
